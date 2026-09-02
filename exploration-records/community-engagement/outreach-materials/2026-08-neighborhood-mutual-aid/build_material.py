from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
ASSET_DIR = BASE_DIR / "assets"
OUTPUT_PATH = BASE_DIR / "近邻互助组_社区合作沟通材料.docx"

# neighborhood_business_proposal -> narrative_proposal tokens
INK = "1B1B1B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "20364D"
MUTED = "666666"
LIGHT_FILL = "F4F6F9"
LIGHT_BLUE = "E8EEF5"
BORDER = "D7DBE2"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM_DXA, start=CELL_SIDE_DXA,
                     bottom=CELL_TOP_BOTTOM_DXA, end=CELL_SIDE_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None,
                 latin="Hiragino Sans GB", east_asia="Hiragino Sans GB"):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=11, color=INK, bold=None, italic=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_text(doc, text, *, size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6,
             line_spacing=1.25, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.3, color=INK)
    return p


def create_numbering_id(doc):
    """Create a real, restartable single-level decimal numbering definition."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.extend([start, num_fmt, lvl_text, lvl_jc])

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "279")
    p_pr.extend([tabs, ind])
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_group(doc, items):
    num_id = create_numbering_id(doc)
    paragraphs = []
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.insert(0, num_pr)
        run = p.add_run(text)
        set_run_font(run, size=10.3, color=INK)
        paragraphs.append(p)
    return paragraphs


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Hiragino Sans GB")
    r_fonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    r_fonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_url(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    label = p.add_run("三个原型在线体验：")
    set_run_font(label, size=10.5, color=MUTED, bold=True)
    add_hyperlink(p, "http://ideal.codeforpeople.cn/", "http://ideal.codeforpeople.cn/")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    return p


def add_section_heading(doc, text):
    p = add_heading(doc, text, 1)
    p.paragraph_format.page_break_before = True
    return p


def add_callout(doc, title, body, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, color=DARK_BLUE, bold=True)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_p.paragraph_format.line_spacing = 1.25
    body_run = body_p.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_picture(doc, filename, alt_text, width_inches=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    inline_shape = run.add_picture(str(ASSET_DIR / filename), width=Inches(width_inches))
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", alt_text)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run("交互原型画面（不接入真实业务数据）")
    set_run_font(cap_run, size=9, color=MUTED, italic=True)


def add_table(doc, headers, rows, widths_dxa, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=DARK_BLUE, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(value)
            set_run_font(run, size=font_size, color=INK, bold=(idx == 0))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instr)
    field_run.append(separate)
    field_run.append(text)
    field_run.append(end)
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Hiragino Sans GB"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
        style.element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Hiragino Sans GB"
        style.font.size = Pt(10.3)
        style.element.rPr.rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
        style.element.rPr.rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.15


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("近邻互助组｜社区合作沟通材料")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("码成仝 · 原型交流版  |  ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(fp)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "近邻互助组｜社区合作沟通材料"
    doc.core_properties.subject = "供布吉街道及相关社区交流的试点讨论稿"
    doc.core_properties.author = "码成仝"
    doc.core_properties.keywords = "近邻互助组, 社区互助, 街坊味, 楼道收一收, 社区试点"

    # Page 1: proposal_centerpiece opening block
    add_text(doc, "社区合作 · 原型交流", size=10.5, color=BLUE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "近邻互助组", size=27, color=NAVY, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line_spacing=1.0)
    add_text(doc, "社区合作沟通材料", size=20, color=NAVY, bold=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line_spacing=1.0)
    add_text(doc, "让邻里间已经发生的信任与互助，更容易被发现、接住和持续",
             size=12.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
             after=20, line_spacing=1.25)

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680])
    set_table_borders(meta, color=BORDER, size="5")
    metadata = [
        ("发起团队", "码成仝"),
        ("沟通对象", "布吉街道及相关社区"),
        ("公司主体", "码成仝"),
        ("版本", "试点讨论稿 · 2026 年 8 月"),
    ]
    for idx, (label, value) in enumerate(metadata):
        cell = meta.cell(idx // 2, idx % 2)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(label + "\n")
        set_run_font(lr, size=9, color=MUTED, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_callout(
        doc,
        "我们希望从一个小场景开始",
        "与社区共同选择一个真实、可控、低风险的问题，用约 2 周完成一次可验证的试点：先解决问题，再根据数据和居民反馈决定是否扩展。",
        fill=LIGHT_BLUE,
    )

    for item in (
        "核心项目是“近邻互助组”：基于真实亲历和既有邻里关系，让消费者之间能够分享有限、可核对的服务经验。",
        "“街坊味”和“楼道收一收”是两个社区场景原型，分别验证供餐协作与楼道回收/保洁协作。",
        "三个页面目前都是交互原型，不接入真实业务数据，也不代表已经上线运营。",
        "建议第一步优先验证“楼道收一收”：范围小、数据少、成效容易观察，适合从一栋楼开始。",
    ):
        add_bullet(doc, item)

    # Page 2: product relationship
    add_section_heading(doc, "一、项目与三个原型的关系")
    add_text(doc, "我们要解决的不是“再建一个群”，而是让已经存在的社区互助不再散落在微信群、私聊和个人记忆中。居民仍然可以在熟悉的关系场求助；只有当有人愿意回应、需要整理信息或承接后续协作时，工具才介入。")

    add_heading(doc, "社区里的三个常见摩擦", 2)
    add_numbered_group(doc, (
        "居民需要维修、保洁、供餐等服务时，很难判断“谁是真的用过”“这条推荐是不是广告”。",
        "热心邻居愿意分享经验，但缺少清楚、可复用、不过度暴露隐私的载体。",
        "一些很小却高频的协作仍靠口头通知和反复确认，物业、服务人员和居民都增加了沟通成本。",
    ))

    add_heading(doc, "三个原型分别回答什么", 2)
    add_table(
        doc,
        ["原型", "整体位置", "解决的问题", "当前适合验证"],
        [
            ["近邻互助组", "核心信任与互助网络", "让“亲自用过”的有限经验帮助邻居作判断", "推荐关系、证据、授权与人机责任"],
            ["街坊味", "社区供餐场景应用", "整理微信里的菜单、预订、配送和线下收款记录", "排菜、接单、履约和对账流程"],
            ["楼道收一收", "楼道协作场景应用", "住户一次扫码告诉师傅哪层需要回收或打扫", "极简上报、楼层待办和一键处理"],
        ],
        [1700, 2100, 3000, 2560],
        font_size=9.1,
    )

    add_heading(doc, "共同原则", 2)
    for item in (
        "真实亲历优先：不做广告排行榜，也不把有限经验包装成平台担保。",
        "人负责：AI 可以整理、匹配和代答，关键承诺仍由真人确认。",
        "数据最少：不抓取微信群，不出售消费记录和关系链。",
        "显式授权：分享、公开、跨应用传递和真人加入都需要用户主动操作。",
        "先验证后扩展：未完成安全、合规与责任确认前，不直接扩大运营。",
    ):
        add_bullet(doc, item)

    # Page 3: neighbor mutual-aid prototype
    add_section_heading(doc, "二、原型一：近邻互助组")
    add_text(doc, "把邻居“亲自用过”的消费与服务经历，变成可核对、可授权、可复用的互助资源。", size=11.5, color=DARK_BLUE, bold=True, after=6)
    add_picture(doc, "近邻互助组.png", "近邻互助组客户体验原型：展示推荐关系、使用次数、最近时间与证据状态", width_inches=5.45)

    add_heading(doc, "一个典型过程", 2)
    add_numbered_group(doc, (
        "居民在原微信群提出真实需求，例如寻找水电维修师傅。",
        "亲自用过的邻居主动回应，并分享一张本人亲历卡。",
        "求助者看到推荐人、使用次数、最近时间、确认状态和本次匹配点。",
        "Agent 先收集现场信息；报价、时间和是否接单由服务者本人确认。",
        "服务完成后，求助者也可以保存自己的经历，未来继续帮助邻居。",
    ))

    add_heading(doc, "社区价值与边界", 2)
    for item in (
        "降低居民选择服务者时的信息不对称和踩坑风险。",
        "让真实经验可复用，又不要求公开完整关系链或私人记忆。",
        "把“谁说的、依据是什么、谁负责承诺”表达清楚。",
        "边界：亲历不是质量保证；无推荐佣金或排名收益；不自动读群；新需求默认独立会话。",
    ):
        add_bullet(doc, item)

    # Page 4: Kith Inn prototype
    add_section_heading(doc, "三、原型二：街坊味")
    add_text(doc, "帮助社区里的小规模供餐者，把每周菜单、邻里预订、配送履约和线下收款记录整理到一条清楚的流程里。", size=11.5, color=DARK_BLUE, bold=True, after=6)
    add_picture(doc, "街坊味.png", "街坊味客户体验原型：顾客微信私聊后由供餐者代录订单并回传订单卡", width_inches=5.45)

    add_heading(doc, "当前原型覆盖", 2)
    for item in (
        "从供餐者已经确认的菜品池中，一次生成并调整五天午、晚餐菜单。",
        "按餐次生成微信分享卡，邻居仍可通过熟悉的微信入口查看和预订。",
        "把顾客自助下单、供餐者代录和既有订餐约定汇总到统一订单中。",
        "显示每餐需要准备多少份、送给谁，并记录配送和线下收款核对状态。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "社区价值与合规边界", 2)
    add_text(doc, "减少依赖群接龙、私聊和个人记忆造成的漏单、漏送、漏记录；让已有的邻里供餐关系更有秩序，而不是建立面向全城的外卖流量平台。", after=2)
    add_text(
        doc,
        "当前仅为流程原型，不代表社区已经开展供餐业务。任何真实试点前，都应先明确食品安全、经营许可、场地和主体责任等条件；在此之前更适合做需求访谈和流程演示。",
        size=9.8,
        color=MUTED,
        after=0,
        line_spacing=1.15,
    )

    # Page 5: recycling/cleaning prototype
    add_section_heading(doc, "四、原型三：楼道收一收")
    add_text(doc, "住户扫一下本层二维码，选择“回收”或“打扫”；师傅看到待处理楼层，处理完成后点一下销单。", size=11.5, color=DARK_BLUE, bold=True, after=6)
    add_picture(doc, "楼道收一收.png", "楼道收一收原型：师傅端楼层面板显示待回收或待打扫楼层")

    add_heading(doc, "为什么适合先试", 2)
    for item in (
        "问题具体：师傅不知道哪层有纸箱、瓶子或需要打扫，只能逐层查看或靠口头通知。",
        "操作极简：二维码已带楼栋和楼层，住户不填表、不选楼层、不必登录。",
        "数据克制：不记录姓名、门牌、手机号、重量、价格、付款、积分或聊天。",
        "结果可测：可以直接观察扫码耗时、漏收情况、处理效率和误操作率。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "建议的 2 周小试点", 2)
    add_table(
        doc,
        ["项目", "建议"],
        [
            ["范围", "1 个小区的 1 栋楼；1 名主要处理人员；约 20–50 户自愿参与"],
            ["准备", "确认放置点；每层张贴唯一二维码；培训师傅使用楼层待办"],
            ["运行", "同楼层提醒合并；处理后点“收好了”；点错可恢复"],
            ["复盘", "第 3 天快速检查；第 14 天共同评估继续、调整或停止"],
        ],
        [1800, 7560],
        font_size=9.7,
    )

    # Page 6+: partnership proposal
    add_section_heading(doc, "五、建议的合作方式")

    stages = [
        ("阶段 1｜现场了解（约半天）", "与社区、物业和实际服务人员确认当前流程、痛点和责任边界；选择一个参与者自愿、可以随时停止的点位。"),
        ("阶段 2｜原型适配与小试点（约 2 周）", "发起团队调整原型、二维码物料和操作说明；社区/物业协助触达参与者；仅收集完成验证所需的最少运行数据和反馈。"),
        ("阶段 3｜共同复盘", "对照预先约定的指标判断是否真正减少沟通成本和漏处理；共同决定继续迭代、换场景验证或停止。"),
    ]
    for title, body in stages:
        add_heading(doc, title, 2)
        add_text(doc, body, after=5)

    add_heading(doc, "建议分工", 2)
    add_table(
        doc,
        ["参与方", "建议承担的工作"],
        [
            ["社区/街道", "指定联系人；协调物业、服务人员和居民访谈；提示政策与公共治理边界；参与复盘"],
            ["发起团队", "需求访谈、原型适配、二维码和说明物料、试点支持、数据最小化、结果整理"],
            ["物业/服务人员", "确认真实工作流程；使用待办和销单；反馈操作困难、漏处理和异常情况"],
            ["居民", "自愿参与；按约定上报；反馈是否易懂、是否愿意继续使用"],
        ],
        [1800, 7560],
        font_size=9.5,
    )

    add_heading(doc, "未经共同确认，本次不做", 2)
    for item in (
        "不公开上线，不向全小区推广，不接入真实收费。",
        "不收集与试点无关的姓名、门牌、手机号或聊天记录。",
        "不抓取微信群内容，不把居民关系用于广告或商业推荐。",
        "不把原型效果当成生产系统的安全、稳定或合规证明。",
        "不在食品安全和经营责任未明确前开展真实供餐运营。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "希望当面沟通的四个问题", 2)
    add_numbered_group(doc, (
        "当前社区最希望降低沟通成本的一个小问题是什么？",
        "“楼道收一收”是否有合适的楼栋、物业或回收/保洁人员愿意一起验证？",
        "社区对居民告知、二维码张贴、数据使用和试点退出有哪些要求？",
        "如果第一轮有效，下一步更适合验证服务亲历，还是其他更迫切的场景？",
    ))
    add_url(doc)
    add_text(doc, "再次说明：全部页面均为交互原型，不接入真实业务数据。", size=9.5, color=MUTED,
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=0)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
